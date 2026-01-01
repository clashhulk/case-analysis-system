"""Export Service for generating DOCX and Markdown files from analysis data"""
from datetime import datetime
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting document analysis to various formats"""

    def generate_docx(
        self,
        filename: str,
        analysis: dict,
        entities: dict,
        extraction: dict
    ) -> bytes:
        """
        Generate a DOCX file from analysis data

        Args:
            filename: Original document filename
            analysis: Analysis results (summary, classification, key_points)
            entities: Extracted entities (people, dates, locations, etc.)
            extraction: Text extraction metadata

        Returns:
            DOCX file as bytes
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"Document Analysis: {filename}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("")

        # Classification & Confidence
        if analysis.get("classification"):
            p = doc.add_paragraph()
            p.add_run("Classification: ").bold = True
            p.add_run(analysis["classification"])

            if analysis.get("confidence"):
                confidence_pct = int(analysis["confidence"] * 100)
                p.add_run(f" ({confidence_pct}% confidence)")

        doc.add_paragraph("")

        # Summary Section
        doc.add_heading("Summary", level=1)
        if analysis.get("summary"):
            doc.add_paragraph(analysis["summary"])
        else:
            doc.add_paragraph("No summary available.")

        # Key Points Section
        if analysis.get("key_points"):
            doc.add_heading("Key Points", level=1)
            for point in analysis["key_points"]:
                doc.add_paragraph(point, style="List Bullet")

        # Entities Section
        doc.add_heading("Extracted Entities", level=1)

        # People
        if entities.get("people"):
            doc.add_heading("People", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Name"
            hdr_cells[1].text = "Role"

            for person in entities["people"]:
                row_cells = table.add_row().cells
                row_cells[0].text = person.get("name", "Unknown")
                row_cells[1].text = person.get("role", "Unknown")

        # Dates
        if entities.get("dates"):
            doc.add_heading("Dates", level=2)
            for date in entities["dates"]:
                doc.add_paragraph(date, style="List Bullet")

        # Locations
        if entities.get("locations"):
            doc.add_heading("Locations", level=2)
            for location in entities["locations"]:
                doc.add_paragraph(location, style="List Bullet")

        # Case Numbers
        if entities.get("case_numbers"):
            doc.add_heading("Case Numbers", level=2)
            for case_num in entities["case_numbers"]:
                doc.add_paragraph(case_num, style="List Bullet")

        # Organizations
        if entities.get("organizations"):
            doc.add_heading("Organizations", level=2)
            for org in entities["organizations"]:
                doc.add_paragraph(org, style="List Bullet")

        # Extraction Metadata
        if extraction:
            doc.add_heading("Extraction Details", level=1)
            p = doc.add_paragraph()
            if extraction.get("extraction_method"):
                p.add_run("Method: ").bold = True
                p.add_run(f"{extraction['extraction_method']}\n")
            if extraction.get("quality_score"):
                p.add_run("Quality Score: ").bold = True
                p.add_run(f"{int(extraction['quality_score'] * 100)}%\n")
            if extraction.get("text_length"):
                p.add_run("Text Length: ").bold = True
                p.add_run(f"{extraction['text_length']} characters\n")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def generate_markdown(
        self,
        filename: str,
        analysis: dict,
        entities: dict,
        extraction: dict
    ) -> str:
        """
        Generate a Markdown file from analysis data

        Args:
            filename: Original document filename
            analysis: Analysis results (summary, classification, key_points)
            entities: Extracted entities (people, dates, locations, etc.)
            extraction: Text extraction metadata

        Returns:
            Markdown content as string
        """
        lines = []

        # Title
        lines.append(f"# Document Analysis: {filename}")
        lines.append("")
        lines.append(f"*Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*")
        lines.append("")

        # Classification
        if analysis.get("classification"):
            classification = analysis["classification"]
            confidence = ""
            if analysis.get("confidence"):
                confidence = f" ({int(analysis['confidence'] * 100)}% confidence)"
            lines.append(f"**Classification:** {classification}{confidence}")
            lines.append("")

        # Summary
        lines.append("## Summary")
        lines.append("")
        if analysis.get("summary"):
            lines.append(analysis["summary"])
        else:
            lines.append("*No summary available.*")
        lines.append("")

        # Key Points
        if analysis.get("key_points"):
            lines.append("## Key Points")
            lines.append("")
            for point in analysis["key_points"]:
                lines.append(f"- {point}")
            lines.append("")

        # Entities
        lines.append("## Extracted Entities")
        lines.append("")

        # People
        if entities.get("people"):
            lines.append("### People")
            lines.append("")
            lines.append("| Name | Role |")
            lines.append("|------|------|")
            for person in entities["people"]:
                name = person.get("name", "Unknown")
                role = person.get("role", "Unknown")
                lines.append(f"| {name} | {role} |")
            lines.append("")

        # Dates
        if entities.get("dates"):
            lines.append("### Dates")
            lines.append("")
            for date in entities["dates"]:
                lines.append(f"- {date}")
            lines.append("")

        # Locations
        if entities.get("locations"):
            lines.append("### Locations")
            lines.append("")
            for location in entities["locations"]:
                lines.append(f"- {location}")
            lines.append("")

        # Case Numbers
        if entities.get("case_numbers"):
            lines.append("### Case Numbers")
            lines.append("")
            for case_num in entities["case_numbers"]:
                lines.append(f"- {case_num}")
            lines.append("")

        # Organizations
        if entities.get("organizations"):
            lines.append("### Organizations")
            lines.append("")
            for org in entities["organizations"]:
                lines.append(f"- {org}")
            lines.append("")

        # Extraction Details
        if extraction:
            lines.append("## Extraction Details")
            lines.append("")
            if extraction.get("extraction_method"):
                lines.append(f"- **Method:** {extraction['extraction_method']}")
            if extraction.get("quality_score"):
                lines.append(f"- **Quality Score:** {int(extraction['quality_score'] * 100)}%")
            if extraction.get("text_length"):
                lines.append(f"- **Text Length:** {extraction['text_length']} characters")
            lines.append("")

        # Footer
        lines.append("---")
        lines.append("*Generated by Case Analysis System*")

        return "\n".join(lines)
