"""Text extraction service for documents

Supports extraction from:
- PDF files using pdfplumber
- Word documents (DOCX) using python-docx
- Scanned images using Tesseract OCR
"""
import re
from pathlib import Path
from typing import Optional
import logging

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None


logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text from various document formats"""

    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize text extraction service

        Args:
            tesseract_path: Path to tesseract executable (for OCR)
        """
        self.tesseract_path = tesseract_path
        if tesseract_path and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

    def extract_from_pdf(self, file_path: str) -> dict:
        """
        Extract text from PDF file using pdfplumber

        Args:
            file_path: Path to PDF file

        Returns:
            dict with text, page_count, method, and metadata
        """
        if not pdfplumber:
            raise ImportError("pdfplumber is required for PDF extraction. Install with: pip install pdfplumber")

        try:
            text_parts = []
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)

            return {
                "text": full_text,
                "text_length": len(full_text),
                "page_count": page_count,
                "method": "pdfplumber",
                "metadata": {
                    "pages_with_text": len(text_parts),
                    "pages_total": page_count
                }
            }

        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def extract_from_docx(self, file_path: str) -> dict:
        """
        Extract text from Word document using python-docx

        Args:
            file_path: Path to DOCX file

        Returns:
            dict with text, paragraph_count, method, and metadata
        """
        if not DocxDocument:
            raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")

        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)

            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\nTables:\n" + "\n".join(table_texts)

            return {
                "text": all_text,
                "text_length": len(all_text),
                "paragraph_count": len(paragraphs),
                "method": "python-docx",
                "metadata": {
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables)
                }
            }

        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def extract_from_image(self, file_path: str) -> dict:
        """
        Extract text from image using Tesseract OCR

        Args:
            file_path: Path to image file (jpg, jpeg, png)

        Returns:
            dict with text, method, and metadata
        """
        if not Image or not pytesseract:
            raise ImportError(
                "PIL and pytesseract are required for OCR. "
                "Install with: pip install Pillow pytesseract"
            )

        try:
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image)

            # Get OCR confidence data
            try:
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in data['conf'] if conf != '-1']
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            except Exception:
                avg_confidence = 0

            return {
                "text": text,
                "text_length": len(text),
                "method": "tesseract-ocr",
                "metadata": {
                    "image_size": image.size,
                    "image_mode": image.mode,
                    "ocr_confidence": round(avg_confidence, 2)
                }
            }

        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from image: {str(e)}")

    def assess_quality(self, text: str) -> float:
        """
        Assess text quality based on various metrics

        Args:
            text: Extracted text to assess

        Returns:
            Quality score between 0 and 1
        """
        if not text or len(text.strip()) < 50:
            return 0.0

        # Metrics for quality assessment
        text_length = len(text)
        word_count = len(text.split())

        # Check for readable characters (letters, numbers, punctuation)
        readable_chars = sum(1 for c in text if c.isalnum() or c.isspace() or c in '.,!?;:-')
        readable_ratio = readable_chars / len(text) if text else 0

        # Check for excessive special characters (indicates poor OCR)
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace() and c not in '.,!?;:-')
        special_ratio = special_chars / len(text) if text else 0

        # Check average word length (gibberish tends to have unusual word lengths)
        avg_word_length = text_length / word_count if word_count > 0 else 0

        # Check for repeated characters (OCR errors often produce these)
        repeated_pattern = len(re.findall(r'(.)\1{3,}', text))
        repeated_ratio = repeated_pattern / len(text) if text else 0

        # Calculate quality score
        quality_score = 0.0

        # Readable ratio should be high (weight: 0.4)
        quality_score += min(readable_ratio, 1.0) * 0.4

        # Special character ratio should be low (weight: 0.2)
        quality_score += max(0, 1 - special_ratio * 5) * 0.2

        # Average word length should be reasonable (3-8 characters) (weight: 0.2)
        if 3 <= avg_word_length <= 8:
            quality_score += 0.2
        elif 2 <= avg_word_length <= 10:
            quality_score += 0.1

        # Repeated character patterns should be minimal (weight: 0.2)
        quality_score += max(0, 1 - repeated_ratio * 20) * 0.2

        return min(round(quality_score, 2), 1.0)

    def extract_text(self, file_path: str, file_type: str) -> dict:
        """
        Main entry point for text extraction

        Routes to appropriate extractor based on file type

        Args:
            file_path: Path to file
            file_type: MIME type or file extension

        Returns:
            dict with extraction results including quality score
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Determine extraction method
        if 'pdf' in file_type.lower() or file_path_obj.suffix.lower() == '.pdf':
            result = self.extract_from_pdf(file_path)

        elif 'word' in file_type.lower() or 'docx' in file_type.lower() or file_path_obj.suffix.lower() in ['.docx', '.doc']:
            result = self.extract_from_docx(file_path)

        elif 'image' in file_type.lower() or file_path_obj.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            result = self.extract_from_image(file_path)

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Assess quality
        quality_score = self.assess_quality(result["text"])
        result["quality_score"] = quality_score

        # Add extraction timestamp
        from datetime import datetime
        result["extracted_at"] = datetime.utcnow().isoformat()

        logger.info(
            f"Text extraction complete: {file_path} | "
            f"Method: {result['method']} | "
            f"Length: {result['text_length']} | "
            f"Quality: {quality_score}"
        )

        return result


# Singleton instance
_text_extraction_service: Optional[TextExtractionService] = None


def get_text_extraction_service(tesseract_path: Optional[str] = None) -> TextExtractionService:
    """
    Get singleton instance of TextExtractionService

    Args:
        tesseract_path: Path to tesseract executable (only used on first call)

    Returns:
        TextExtractionService instance
    """
    global _text_extraction_service

    if _text_extraction_service is None:
        _text_extraction_service = TextExtractionService(tesseract_path=tesseract_path)

    return _text_extraction_service
